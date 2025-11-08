"""
AI Resume Transformation FastAPI Application
Complete, self-contained application for transforming resumes using Google Gemini API.
"""

# ============================================================================
# PART 1: IMPORTS
# ============================================================================

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import httpx
import asyncio
import json
import re
import os
import io
import zipfile
import aiofiles
import logging
from typing import Optional, List, Dict, Any
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black, white, HexColor, Color
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ============================================================================
# LOGGING CONFIGURATION
# ============================================================================

# Create logs directory if it doesn't exist
# os.makedirs("logs", exist_ok=True)

# Configure logging with both file and console output
log_filename = f"logs/resume_transform_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

# logging.basicConfig(
#     level=logging.DEBUG,
#     format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
#     handlers=[
#         logging.FileHandler(log_filename),
#         logging.StreamHandler()
#     ]
# )
# Configure logging for console output (Render-friendly)
logging.basicConfig(
level=logging.INFO, # Use INFO for production, DEBUG is too noisy
format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
handlers=[
logging.StreamHandler()
]
)

logger = logging.getLogger(__name__)
logger.info("=" * 80)
logger.info("AI Resume Transformation Service Started")
logger.info(f"Log file: {log_filename}")
logger.info("=" * 80)


# ============================================================================
# PART 2: HELPER FUNCTIONS AND CONSTANTS
# ============================================================================

# Prohibited clichés list for filtering
PROHIBITED_CLICHES = [
    "action and results oriented", "team player", "a quick and agile learner",
    "self-motivated", "innovative thinker", "go-getter", "detail-oriented",
    "results-driven", "passionate professional", "hard worker", 
    "think outside the box", "proactive", "willingness to take initiative",
    "process driven"
]

def filter_cliches(skills_string: str) -> str:
    """Deterministically removes clichés from an LLM-generated skill list
       using an exact match."""
    if not skills_string:
        return ""
    
    # Create a set of lowercase, stripped clichés for fast, exact matching
    cliche_set = {c.lower().strip() for c in PROHIBITED_CLICHES}
    
    skills_list = [s.strip() for s in skills_string.split(',')]
    filtered = []
    
    for skill in skills_list:
        skill_lower_stripped = skill.lower().strip()
        # Only add the skill if it's not empty and not an exact match
        # for a cliché
        if skill_lower_stripped and skill_lower_stripped not in cliche_set:
            filtered.append(skill.strip()) # Append the original, cased skill
            
    return ', '.join(filtered)

def enforce_absolute_urls(contact_line: str) -> str:
    """Ensures all links start with https:// and cleans duplicates."""
    if not contact_line:
        return ""
    
    # First, fix any http://
    contact_line = contact_line.replace('http://', 'https://')
    
    # CRITICAL: Fix duplicates this bug created
    contact_line = contact_line.replace('https://https://', 'https://')
    
    # Now, add https:// *only if* the protocol is missing
    if 'linkedin.com' in contact_line and 'https://linkedin.com' not in contact_line:
        contact_line = contact_line.replace('linkedin.com', 'https://linkedin.com')
    if 'github.com' in contact_line and 'https://github.com' not in contact_line:
        contact_line = contact_line.replace('github.com', 'https://github.com')
    if 'netlify.app' in contact_line and 'https://' not in contact_line:
         contact_line = contact_line.replace('ayaanahmad-portfolio.netlify.app', 'https://ayaanahmad-portfolio.netlify.app')
    
    # Final cleanup of any duplicates the previous logic might have missed
    contact_line = contact_line.replace('https://https://', 'https://')
    return contact_line


# ============================================================================
# PART 3: FASTAPI APP SETUP
# ============================================================================

app = FastAPI(title="AI Resume Transformer")

# ============================================================================
# CORS CONFIGURATION (for frontend development on port 5500)
# ============================================================================

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5500",
        "http://127.0.0.1:5500",
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:8080",
        "http://127.0.0.1:8080",
        "*"  # Allow all origins (for development only)
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

logger.info("✓ CORS middleware configured for frontend development")

# --- Gemini API Configuration ---
# --- Gemini API Configuration ---
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"

# Securely load API key from environment variables
API_KEY = os.environ.get("GEMINI_API_KEY")

# Add a check to fail fast if the key is missing
if not API_KEY:
    logger.critical("FATAL ERROR: GEMINI_API_KEY environment variable is not set.")
    # This will stop the app from starting if the key is missing
    raise ValueError("GEMINI_API_KEY is not set. The application cannot start.")

logger.info(f"Gemini API URL configured: {GEMINI_API_URL}")
logger.info(f"API Key present: {bool(API_KEY)}")


# ============================================================================
# PART 4: MASTER LLM PROMPT TEMPLATE
# ============================================================================

RESUME_TRANSFORM_PROMPT_TEMPLATE = """Comprehensive Resume Transformation Prompt

# ROLE
You are an elite ATS Resume Optimization Specialist with deep expertise in Taleo, Greenhouse, Human Review, and Jobscan scoring algorithms. Your mission is to transform the provided resume to achieve maximum scores across all four systems while maintaining authenticity, natural language, and **fitting to 1 page (<=370 words)**.
---
# SCORING SYSTEMS (RESEARCH-BASED)
## **OVERALL SCORING FORMULA**
- **Final Score** = (30% × Taleo) + (25% × Greenhouse) + (25% × Human Review) + (20% × Jobscan)
- **Target**: 90+ overall score; 100% Jobscan score with 0 issues
- **Critical Thresholds**: Taleo hard skills ≥50% to avoid penalties; Jobscan hard skills 100% presence with frequency >= JD counts; total word count <=370 for 1-page
---
## **SYSTEM 1: TALEO SCORER (30% Weight)**
**Type**: Gen1 ATS - Exact keyword matching + format compliance
### **Formula Components:**

Taleo Score = (
    50% × Hard Skills Coverage (EXACT match) +
    20% × Job Title Match +
    15% × Soft Skills Coverage (stemmed match) +
    10% × Education Match +
    5% × Other Keywords Coverage
) × 100
CRITICAL PENALTY: If Hard Skills Coverage < 50%, apply penalty:
Final Score = Score × (1 - ((0.5 - coverage) × 100) / 100)

### **Detailed Scoring Logic:**
**1. Hard Skills Coverage (50% - EXACT MATCH ONLY)**
- Extract all tier-1 technical skills and phrases from job description (include multi-word like "data visualization")
- Match EXACT keywords (case-insensitive) from resume
- Formula: `matches / total_tier1_skills`
- **CRITICAL**: Must achieve 50%+ or severe penalty applies
**2. Job Title Match (20%)**
- Compare most recent job title with JD domain/title
- Split titles into words, calculate word overlap
- Formula: `min(overlapping_words / jd_domain_words, 1.0)`
**3. Soft Skills Coverage (15% - STEMMED MATCH)**
- Identify soft skills from tier-2 keywords (communication, leadership, teamwork, problem-solving, collaboration, management, presentation)
- Use word stemming for matching (e.g., "analyze" matches "analysis")
- Formula: `matches / total_soft_skills`
**4. Education Match (10%)**
- Check for Bachelor's/Master's degree
- Match to seniority level expectation
- Scoring:
  - Has degree + appropriate seniority = 1.0
  - Has degree but mismatched = 0.7
  - No degree = 0.5
**5. Other Keywords (5%)**
- Tier-2 keywords excluding soft skills
- Exact match ratio
---
## **SYSTEM 2: GREENHOUSE SCORER (25% Weight)**
**Type**: Gen2 ATS - Format compliance + structure + parsability
### **Formula:**

Greenhouse Score = (
    50% × Format Compliance +
    30% × Section Structure +
    20% × Data Parsability
) × 100

### **Detailed Scoring Logic:**
**1. Format Compliance (50%)**
- **INSTANT FAIL (0.0)**: Tables present
- **Penalties**:
  - Columns: -0.3
  - Graphics/images: -0.2
  - Missing contact info: -0.3
- Formula: `max(1.0 - total_penalties, 0.0)`
**2. Section Structure (30%)**
- **Required sections** (each missing = -0.3):
  - Experience
  - Education
- **Recommended sections** (each missing = -0.1):
  - Skills
- Formula: `max(1.0 - total_deductions, 0.0)`
**3. Data Parsability (20%)**
- **Date format validation** (invalid = -0.15 per experience):
  - Valid: "Month YYYY - Month YYYY" or "Month YYYY - Present"
- **Education year** (missing/invalid = -0.1 per degree):
  - Must be 4-digit year
- **Link Validation:** (invalid = -0.1 per link):
  - Valid: `https://linkedin.com/in/user` ✓
  - Invalid: `linkedin.com/in/user` or `/in/user` ✗
- Formula: `max(1.0 - total_deductions, 0.0)`
---
## **SYSTEM 3: HUMAN REVIEW SCORER (25% Weight)**
**Type**: Deductive scoring - starts at 100, subtracts flaws
### **Formula:**

Start: 100.0
Deduct:
  - Missing quantification: min(unquantified_bullets × 2.0, 20.0)
  - Weak action verbs: min(weak_verb_bullets × 1.0, 10.0)
  - Missing CAR format: min(missing_CAR_bullets × 10.0, 30.0)
  - Length issues: -20.0 if word_count > 370 (1-page cap)
Final: max(score, 0.0)

### **Detailed Scoring Logic:**
**1. Quantification Check (max -20 points)**
- **Every bullet must contain numbers**: percentages, dollar amounts, counts, metrics. **A minimum of 5 total quantified metrics** must be present in the resume.
- Detection: Regex `\d+[\d,]*[%$]?|\$\d+`
- Penalty: 2 points per unquantified bullet (capped at 20)
**2. Weak Action Verbs (max -10 points)**
- **Prohibited first words**: helped, assisted, supported, contributed, worked on
- Penalty: 1 point per weak verb (capped at 10)
**3. CAR Format - Context-Action-Result (max -30 points)**
- **Every bullet must show results/impact**
- Detection: Must contain result indicators:
  - "resulting in", "led to", "improved", "increased", "reduced", "achieved", "enabled", "delivered"
- Penalty: 10 points per bullet without CAR (capped at 30)
**4. Length Check (max -20 points)**
- **CRITICAL**: Penalty: -20 if word_count > 370 OR if content spills to a second page.
---
## **SYSTEM 4: JOBSCAN SCORER (20% Weight)**
**Type**: Keyword-focused matcher - Presence, frequency, and ATS/readability checks
### **Formula:**

Jobscan Score = (
    50% × Hard Skills Match (presence + frequency) +
    20% × Soft Skills Match +
    15% × Searchability Compliance +
    15% × Recruiter Tips Alignment
) × 100

### **Detailed Scoring Logic:**
**1. Hard Skills Match (50%)**
- Extract all technical skills/phrases from JD (single/multi-word)
- Require 100% presence; frequency in resume >= JD's count for each (e.g., "product analytics" >=3)
- Formula: `(present_keywords / total_jd_keywords) + min((resume_freq / jd_freq), 1.0 per keyword)`
- Penalty for misses/low freq: -10% per missing/low keyword (capped at 4 issues)
**2. Soft Skills Match (20%)**
- Extract behavioral phrases (e.g., "problem-solving skills")
- Stemmed match with frequency >= JD (e.g., "problem-solving skills" >=1)
- Formula: `matches / total_soft_skills`
**3. Searchability Compliance (15%)**
- Check contact (address required in standard format, e.g., "City, State, Country" **unlabeled**)
- Penalties: -0.5 for no/poorly parsed address (e.g., using a label like "Address:"); -0.2 for special filename chars.
- Formula: `max(1.0 - total_penalties, 0.0)`
**4. Recruiter Tips Alignment (15%)**
- Check measurable results (5+), positive tone, job level match (e.g., 3+ "Entry-Level New Grad" mentions)
- Penalty: -10 for warnings (e.g., no explicit alignment phrasing)
---
# TRANSFORMATION WORKFLOW
## **PHASE 1: ANALYSIS**
1. Extract all tier-1 skills/phrases (hard skills including multi-word like "product analytics", "business intelligence", "R", "Power BI", "Mixpanel"; education fields like "Computer Science", "mathematics", "statistics")
2. Extract all tier-2 skills (soft skills + other keywords/phrases, e.g., "data-driven", "key metrics", "strong analytical and problem-solving skills")
3. Identify job domain/title keywords and phrases (e.g., "Product Analyst", "New Grad / Entry Level")
4. Determine seniority level expectation and frequency counts from JD for all keywords (include bonus points)
5. Calculate baseline scores for original resume (include Jobscan with frequency breakdowns)
6. Calculate effective project development time: If provided, use TIME_IN_WEEKS x AI_MULTIPLIER (default TIME_IN_WEEKS=2, AI_MULTIPLIER=2 if not provided; valid multipliers: 2, 3, 5). Assume 20 hours/week effort.
7. Prioritize content: Select top 2 relevant work ex, top 2 education, regenerate 2 projects.
## **PHASE 2: PRESERVATION**
**DO NOT MODIFY:**
- Work experience: company names, job titles, employment dates
- Education: degree names, institution names, graduation years
- Core factual responsibilities that are verifiable
## **PHASE 3: TRANSFORMATION**
### **Content Selection for 1-Page (HARD RULES)**
- **Limit to**: Contact, Summary (**3-4 lines MAX**), Experience (top 2 relevant entries, **2 bullets MAX**), Projects (2 regenerated, **2-3 bullets MAX**), Skills (compressed), Education (top 2 entries)
- **Word Count**: Total word count **MUST** be <=370.
- **Compression**: **RUTHLESSLY edit all bullets to 1 line (approx. 15-20 words). 2 lines is an exception only if critical keywords are needed.**
### **A. Projects Section (COMPLETE REGENERATION)**
Generate 2 NEW projects that:
**Requirements:**
1. **Keyword saturation**: Include ALL tier-1 hard skills/phrases from JD (including bonus points) with frequency >= JD's (e.g., "product analytics" 3+ times, "R" 1+, "Power BI" 1+, "Mixpanel" 1+)
2. **Domain alignment**: Projects must be relevant to target role
3. **Quantification**: Every bullet must have 1-2+ metrics (ensure 5+ total for Jobscan)
4. **CAR format**: Every bullet must show measurable results
5. **Credibility**: Projects must sound realistic and achievable; scale metrics to fit calculated time (e.g., 10K-100K data points/users for 40-120 hours, adjustable for seniority)
6. **Technical depth**: Use exact terminology/phrases from JD
7. **RUTHLESS CONCISENESS**: Bullets **MUST be 1 line (approx. 15-20 words) MAXIMUM.**
**Project Structure Template:**

PROJECT NAME | Technologies: [exact tier-1 keywords/phrases from JD]
Duration: Month YYYY - Month YYYY

[Strong action verb] [technical task] using [tier-1 phrase], resulting in [quantified outcome] (1 line MAX, approx 15-20 words)

[Strong action verb] [technical task] using [tier-1 phrase], resulting in [quantified outcome] (1 line MAX, approx 15-20 words)

### **B. Experience Section (OPTIMIZE BULLETS)**
**For each selected work experience (top 2):**
1. **Inject tier-1 keywords/phrases naturally** (ensure frequency >= JD across resume).
2. **Add quantification**: (e.g., "Optimized system performance, reducing API response time by 45%")
3. **Ensure CAR format**: (e.g., "Led AWS cloud migration... resulting in 60% cost reduction")
4. **Replace weak verbs**: (e.g., helped → Led; assisted → Engineered; worked on → Architected)
5. **Maintain authenticity**: Stay within reasonable bounds; limit to **2 bullets MAXIMUM**.
6. **Compress for Density**: Every bullet **MUST be 1 line (approx. 15-20 words) MAXIMUM.**
### **C. Skills Section (EXACT KEYWORD MATCHING & HYPER-COMPRESSED)**
**Structure:**
SKILLS
[One single, comma-separated list of ALL tier-1 and tier-2 skills/phrases. e.g., SQL, Python, R, Statistics, Mathematics, Tableau, Power BI, product analytics, data interpretation, Strong analytical and problem-solving skills, Communication]

**Critical Rules:**
- Use EXACT keyword strings/phrases from JD (case can vary, include multi-word).
- **CRITICAL: Consolidate ALL skills into a single, comma-separated list under the "SKILLS" header. DO NOT use sub-headings (e.g., "Technical Skills:").** This is a deliberate compression strategy to save vertical space and guarantee the 1-page fit.
- The single list must be 1-2 lines maximum.
### **D. Format Optimization (FOR AI OUTPUT)**
**Required:**
- **Output plain text only.**
- Clean, single-column layout
- Standard section headers: Professional Summary, Experience, Projects, Skills, Education
- Bullet points (`-` or `*`) for all descriptions (**1 line MAX, approx 15-20 words**)
- Consistent date format: "Month YYYY - Month YYYY" or "Month YYYY - Present"
- No tables, no columns, no graphics, no special characters.
- Emphasize job level in Summary multiple times (e.g., "Entry-Level New Grad... aligning with new grad responsibilities") and add company tip (e.g., "[Company from JD] ([website from JD])"). Summary MUST be **3-4 lines MAXIMUM.**
- **Prohibited Clichés:** DO NOT use common resume clichés in the Professional Summary or anywhere else. Banned phrases include: "Action and results oriented", "Team player", "A quick and agile learner", "Self-motivated", "Innovative thinker", "Go-getter", "Detail-oriented", "Results-driven", "Passionate professional", "Hard worker", "Think outside the box". Instead, demonstrate these qualities through quantified achievements in your bullet points.
- **Contact Info Format**: The [Full Name] must be on its own line at the top. Follow with a 3-line header. **CRITICAL**: Links MUST be full absolute URLs (e.g., `https://...`) and the address MUST NOT have a label (e.g., "Address:").
  - Line 1: [Full Name]
  - Line 2: [email@domain.com] | [Phone] | [https://linkedin.com/in/user] | [https://github.com/user] | [https://portfolio.url]
  - Line 3: [City, State, Country]
- **Education Format**: Format each entry over 2 lines, matching the provided .docx example:
  - Line 1: [Degree Name] ([Optional Specialization])
  - Line 2: [Institution Name], [Location] [Start Month YYYY] - [End Month YYYY]
**Sections Order:**
1. Contact Information (Name + 3-line format)
2. Professional Summary (3-4 lines, keyword-rich)
3. Experience (top 2 entries, 2 bullets each)
4. Projects (2 entries, 2-3 bullets each)
5. Skills (Single-list, hyper-compressed format)
6. Education (top 2 entries, 2-line format each)
---
# INPUT REQUIREMENTS
**Provide:**
1. [ORIGINAL RESUME - attached]
2. [JOB DESCRIPTION - pasted here]
3. [TIME_IN_WEEKS - integer, e.g., 2 for 2 weeks prep time; optional, default 2 if not provided]
4. [AI_MULTIPLIER - integer, e.g., 3 for x3; optional, choices 2, 3, 5; optional, default 2 if not provided]
---
# OUTPUT FORMAT
## **Part 1: Transformation JSON**
[Provide the entire output as a single, valid JSON object. Do not include any text before or after the JSON block. Populate the `resumeObject` with the transformed resume content, adhering to the structure defined below. **CRITICAL**: All strings in the JSON, especially URLs, must be plain text without any Markdown formatting.]

Here is the resume to transform:

[Paste your resume here]

Here is the job description:

[Paste job description here]

Time in weeks for project development: [Optional: TIME_IN_WEEKS e.g., 2]

AI Multiplier for project scaling: [Optional: AI_MULTIPLIER e.g., 3]

Output the following JSON structure:

```json
{
  "initialScores": {
    "overallScore": 0,
    "kudoswallPredicted": 0,
    "taleo": {
      "score": 0,
      "hardSkillsCoverage": "X% (matched Y/Z)",
      "jobTitleMatch": "X%",
      "softSkills": "X%",
      "education": "X/1.0",
      "otherKeywords": "X%"
    },
    "greenhouse": {
      "score": 0,
      "formatCompliance": "X/1.0",
      "sectionStructure": "X/1.0",
      "parsability": "X/1.0"
    },
    "humanReview": {
      "score": 0,
      "quantifiedBullets": "Y/Z",
      "weakVerbs": 0,
      "carFormat": "Y/Z",
      "wordCount": 0,
      "penalty": "X"
    },
    "jobscan": {
      "score": 0,
      "hardSkillsMatch": "X% (missed Y issues)",
      "softSkillsMatch": "X%",
      "searchability": "X/1.0",
      "recruiterTips": "X/1.0"
    }
  },
  "transformedResume": {
    "suggestedFilename": "[FirstLast]_[JobTitle]_Resume.pdf",
    "resumeObject": {
      "contactInfo": {
        "name": "[Full Name]",
        "contactLine": "[email@domain.com] | [Phone] | [https://linkedin.com/in/user] | [https://github.com/user] | [https://portfolio.url]",
        "location": "[City, State, Country]"
      },
      "professionalSummary": "[String of the full summary...]",
      "experience": [
        {
          "title": "[Job Title]",
          "companyLocation": "[Company Name], [Location]",
          "dates": "[Month YYYY] - [Month YYYY]",
          "bullets": [
            "[Bullet 1... (1 line MAX, approx 15-20 words)]",
            "[Bullet 2... (1 line MAX, approx 15-20 words)]"
          ]
        }
      ],
      "projects": [
        {
          "titleAndTech": "[Project Name] | Technologies: [Keywords...]",
          "dates": "[Month YYYY] - [Month YYYY]",
          "bullets": [
            "[Bullet 1... (1 line MAX, approx 15-20 words)]",
            "[Bullet 2... (1 line MAX, approx 15-20 words)]"
          ]
        }
      ],
      "skills": "[One single, comma-separated string of ALL skills/phrases. e.g., SQL, Python, R, Statistics, Tableau, product analytics, Strong analytical and problem-solving skills]",
      "education": [
        {
          "degree": "[Degree Name] ([Optional Specialization])",
          "institutionAndDates": "[Institution Name], [Location] [Start Month YYYY] - [End Month YYYY]"
        }
      ]
    }
  },
  "finalScores": {
    "overallScore": 0,
    "kudoswallPredicted": 0,
    "taleo": {
      "score": 0,
      "improvement": "+X",
      "hardSkillsCoverage": "X% (matched Y/Z)",
      "jobTitleMatch": "X%",
      "softSkills": "X%",
      "education": "X/1.0",
      "otherKeywords": "X%"
    },
    "greenhouse": {
      "score": 0,
      "improvement": "+X",
      "formatCompliance": "1.0",
      "sectionStructure": "1.0",
      "parsability": "1.0"
    },
    "humanReview": {
      "score": 0,
      "improvement": "+X",
      "quantifiedBullets": "Y/Z",
      "weakVerbs": 0,
      "carFormat": "Y/Z",
      "wordCount": 0,
      "penalty": "0"
    },
    "jobscan": {
      "score": 0,
      "improvement": "+X",
      "hardSkillsMatch": "100% (0 issues)",
      "softSkillsMatch": "100%",
      "searchability": "1.0/1.0",
      "recruiterTips": "1.0"
    }
  },
  "transformationSummary": {
    "keywordsAdded": [
      "keyword (freq)",
      "multi-word phrase (freq)"
    ],
    "selectedContent": "Top 2 work ex (2 bullets each), 2 projects (2-3 bullets each), top 2 education",
    "bulletsModified": 0,
    "quantificationAdded": 0,
    "weakVerbsEliminated": 0,
    "carFormatApplied": 0,
    "formatViolationsFixed": [
      "list of fixes",
      "e.g., absolute https:// links",
      "e.g., unlabeled address",
      "e.g., compressed skills section"
    ],
    "jobscanIssuesResolved": [
      "list of all fixed issues"
    ],
    "onePageCompliance": "Word count [X] <=370. 1-Page fit confirmed.",
    "effectiveProjectTime": "[TIME_IN_WEEKS] weeks x [AI_MULTIPLIER] = [calculated hours] hours at 20 hours/week"
  }
}
```
"""


# ============================================================================
# PART 4B: NEW 3-CALL PIPELINE PROMPTS
# ============================================================================

JD_ANALYSIS_PROMPT_TEMPLATE = """You are a meticulous Job Description Analyzer. Your sole task is to scan the provided
job description and extract all critical requirements into a structured JSON object.

# Exclusion Rules (CRITICAL)
- **DO NOT** extract geographic locations (e.g., "Noida", "Bangalore", "Ahmedabad", "Chennai", "Gurgaon").
- **DO NOT** extract company-specific HR or hiring process terms (e.g., "Interview rounds", "Resume based shortlisting", "Online Aptitude Test").
- **DO NOT** extract generic academic requirements (e.g., "7+ CGPA", "70%+", "No Backlog", "Strong academic performance").
- **DO NOT** extract vague, non-skill phrases (e.g., "work with supervisor", "complete the given task", "adhering to quality", "willingness to take initiative", "cultural fit").
- **DO NOT** extract generic company roles (e.g., "Business Analysts", "senior team members").
- **DO NOT** extract common cliches (e.g., "Action and results oriented", "Team player", "A quick and agile learner", "Self-motivated", "Innovative thinker", "Go-getter", "Detail-oriented", "Results-driven", "Passionate professional", "Hard worker", "Think outside the box").
- **FOCUS ONLY** on transferable, measurable skills, tools, and domain concepts (e.g., "SQL", "Tableau", "data engineering", "analytical skills").

# Job Description
[Insert Job Description Here]

# Output Format
Return **only** a valid JSON object in the following format.
Do not invent requirements not explicitly mentioned.

{
  "company_name": "(e.g., Zomato, Google, Microsoft)",
  "keywords": {
    "hard_skills": ["(e.g., SQL, Python, Tableau)"],
    "soft_skills": ["(e.g., communication, analytical)"],
    "tools": ["(e.g., Mixpanel, Firebase, Power BI)"],
    "domain_phrases": ["(e.g., product analytics, business intelligence, A/B Testing)"]
  },
  "explicit_requirements": {
    "education": ["(e.g., Bachelor's, Master's, Computer Science)"],
    "experience_level": "(e.g., 3+ years, Entry-Level, New Grad)"
  },
  "metric_indicators": [
    "(e.g., proven track record, data-driven, impact key metrics)"
  ]
}
"""

KEYWORD_VERIFICATION_PROMPT_TEMPLATE = """You are a meticulous ATS Keyword Verifier.
Your sole task is to find keywords from the Job Description that are MISSING from the provided Keyword List.

# Job Description
[Insert Job Description Here]

# Current Keyword List (JSON)
[Insert Keyword List JSON Here]

# Instructions
1.  Read the Job Description exhaustively.
2.  Read the Current Keyword List.
3.  Identify *every single* skill, tool, or phrase (hard or soft) from the Job Description that is NOT present in the Current Keyword List.
4.  Pay special attention to multi-word phrases (e.g., "technical competence", "business requirements").

# Output Format
Return **only** a valid JSON object containing a single key "missing_keywords" with a list of the missing strings.
If none are missing, return an empty list.

{
  "missing_keywords": [
    "(e.g., business requirements)",
    "(e.g., quality standards)",
    "(e.g., technical competence)"
  ]
}
"""

DYNAMIC_VERIFY_AND_CORRECT_PROMPT_TEMPLATE = """You are an elite ATS Resume QA Specialist. Your job is to surgically correct a
draft resume to ensure it 100% aligns with a dynamic "JD Analysis Report".
This is the final quality gate. You must be merciless.

# Exclusion Rules (CRITICAL)
- **DO NOT** extract or append geographic locations (e.g., "Noida", "Bangalore").
- **DO NOT** extract or append HR/hiring terms (e.g., "Interview rounds", "Resume based shortlisting").
- **DO NOT** extract or append generic academic requirements (e.g., "7+ CGPA", "No Backlog").
- **DO NOT** extract or append vague, non-skill phrases (e.g., "work with supervisor", "complete the given task").
- **DO NOT** extract or append common cliches (e.g., "Action and results oriented", "Team player", "A quick and agile learner", "Self-motivated", "Innovative thinker", "Go-getter", "Detail-oriented", "Results-driven", "Passionate professional", "Hard worker", "Think outside the box").
- **FOCUS ONLY** on transferable, measurable skills, tools, and domain concepts.

# TASK HIERARCHY (CRITICAL)
Your tasks MUST be performed in this exact order of priority.
If a lower-priority constraint (like Word Count) prevents you
from completing a higher-priority task (like Keyword Inclusion),
you MUST violate the lower-priority constraint to succeed at the higher one.

---
**1. PRIORITY #1: Core Content & Metrics (NON-NEGOTIABLE)**
   - **Company Name:** The 'professionalSummary' MUST contain the `company_name` value found in the `jdAnalysisReport`.
   - **Address Pincode:** The 'location' field MUST be complete. It must be in the format "[City, State, Pincode]" (e.g., "New Delhi, India, 110025").
   - **Action:** Check if the 'location' string contains a 6-digit pincode.
   - **Action (Fix):** If no pincode is present, find the location of the *most recent* (first) entry in the 'education' section (e.g., "New Delhi" from "Jamia Millia Islamia, New Delhi"). Use a valid, standard pincode for that city (e.g., 110025 for New Delhi) and append it to the location string.
   - **Action (No Label):** Ensure the location string does NOT start with a label like "Address: ".
   - **Hard Requirements:** The 'professionalSummary' or 'education' MUST
     match the JD's `explicit_requirements`.
   - **Metrics Requirement:** If `metric_indicators` is not empty, the resume MUST
     contain 5 or more unique metrics (e.g., %, $, 1M+, 5+).

**2. PRIORITY #2: Formatting & Length (VIOLATE IF NEEDED)**
   - **Rule:** These are targets. They are less important than Priority #3 or #4.
   - **Professional Summary:** Target <= 40 words.
   - **Bullet Length:** Target <= 1 line (approx 15-20 words).
   - **Total Word Count:** Target <= 380 words.
   - **Action:** You MUST obey these limits *unless* it is impossible to
     also satisfy Priority #1, #3, or #4.

**3. PRIORITY #3: Keyword Weaving (Best Effort)**
   - **Rule:** Attempt to naturally weave all **valid** keywords (hard, soft, tools, phrases)
     from the report into the 'professionalSummary', 'experience', or 'projects' bullets,
     **while obeying the Exclusion Rules.**
   - **Rule (Frequency):** Attempt to meet the frequency requirement for **valid** keywords
     (like 'communication').
   - **Action:** This is a "best effort" attempt. Do not sacrifice readability.

**4. PRIORITY #4: FINAL MANDATE (NON-NEGOTIABLE)**
   - **Rule:** This task is the absolute highest priority and must be
     executed last, after all other edits.
   - **Action:** Scan the *entire* resume one final time against the
     `jdAnalysisReport`'s keyword list, **while strictly obeying the Exclusion Rules above.**
   - **Action (Fix):** For ANY and ALL **valid** keywords (like 'Data Engineering') or
     frequency requirements (like a second 'communication') that are *still*
     missing, you **MUST** add them to the `skills` string.
   - **Action (Exact Match):** This check **MUST** be an **EXACT**, case-insensitive
     string match. **Do not** skip a keyword like 'technical skills' just
     because 'technical competence' is already present. If the report lists a
     keyword, and it's not an *exact* match in the resume, it **MUST** be
     appended.
   - **Action (Brute-Force):** Append the missing **valid** words/phrases to the end of
     the `skills` string. This is not optional. 100% **valid** keyword inclusion
     is the only acceptable output, even if it makes the skills string
     long or violates the 380-word target.

---
# CONFLICT RULES

**1. Project Handling:**
   - **Rule:** The 'projects' section is *generated content*.
   - **Action:** **DO NOT** fact-check the project content against the
     `explicit_requirements`.
   - **Your only job** for projects is to ensure they help satisfy
     Priority #1 (Metrics) and Priority #3/#4 (Keywords).

**2. Summary vs. Keywords:**
   - **Rule:** You must *try* to meet the 40-word summary target.
   - **Action:** If you must remove a **valid** keyword from the 'professionalSummary'
     to meet this target, you **MUST** ensure that same keyword is
     present elsewhere (e.g., in the `skills` list, per Priority #4).
     **Keyword preservation (Priority #4) is always more important.**

---

# JD ANALYSIS REPORT (JSON)
[Insert jdAnalysisReport JSON Here]

# INPUT DRAFT RESUME (JSON)
[Insert resumeObject JSON string here]

# OUTPUT
Return **only** the complete, corrected `resumeObject` JSON. Do not write
any explanation.
"""

FINAL_BRUTEFORCE_QA_PROMPT = """You are an elite ATS Resume QA Specialist performing the FINAL check.
Your primary job is 100% keyword coverage via brute-force append.
Secondary jobs are summary fixing and word count trimming.

# TASK HIERARCHY (CRITICAL - Highest to Lowest)

1.  **KEYWORDS (BRUTE-FORCE APPEND - NON-NEGOTIABLE):**
    * Compare the "REQUIRED KEYWORDS" list to the "CURRENT RESUME KEYWORDS" list.
    * Identify ANY keywords from the "REQUIRED" list that are MISSING from the "CURRENT" list (case-insensitive check).
    * You **MUST** append every single missing keyword to the end of the
        `skills` string within the `resumeObject_draft`. Do not try to weave them in. Just append.
        100% keyword coverage in the final `skills` list is mandatory.

2.  **SUMMARY (NON-NEGOTIABLE):**
    * Ensure `professionalSummary` contains `target_job_title`.
    * Ensure `professionalSummary` contains `company_name`.
    * Rewrite the summary ONLY IF NEEDED to meet these, keeping it **under 40 words**.

3.  **WORD COUNT (TRIM IF NEEDED):**
    * Check total word count of the *modified* resume. If > 370 words, you MUST
        shorten bullets in the `projects` section until count is <= 370.
        Prioritize keeping keywords/metrics when trimming.

# REQUIRED KEYWORDS (from JD Analysis)
[Insert jd_all_skills_str Here]

# CURRENT RESUME KEYWORDS (from Draft Resume)
[Insert resume_skills_str Here]

# TARGET JOB TITLE
[Insert Target Job Title String Here]

# COMPANY NAME
[Insert Company Name String Here]

# INPUT DRAFT RESUME (JSON)
[Insert resumeObject JSON string here]

# OUTPUT (JSON Mode)
Return **only** the complete, corrected `resumeObject` JSON including the
updated skills string. Ensure it's valid JSON. Do not write any explanation.
"""


# ============================================================================
# PART 5: PDF PARSING HELPERS (PyMuPDF)
# ============================================================================

def find_urls_in_text(text: str) -> List[str]:
    """Finds URLs written as plain text."""
    url_regex = r'((?:https?://|www\.)[a-zA-Z0-9./?=&-_]+)'
    return re.findall(url_regex, text)


def identify_link_type(url: str) -> str:
    """Identifies the type of link."""
    url_lower = url.lower()
    if 'linkedin.com' in url_lower:
        return 'LinkedIn'
    if 'github.com' in url_lower:
        return 'GitHub'
    if 'dev' in url_lower or 'portfolio' in url_lower or 'behance' in url_lower:
        return 'Portfolio'
    return 'Other'


def parse_resume_pdf(pdf_path: str) -> Optional[str]:
    """
    Parses a PDF resume to extract its full text content and embed links.
    Returns a single string of the resume's text with embedded links.
    """
    logger.info(f"📄 Starting PDF parsing: {pdf_path}")
    full_text = ""
    link_details = []

    try:
        logger.debug(f"Opening PDF file: {pdf_path}")
        doc = fitz.open(pdf_path)
        logger.info(f"✅ PDF opened successfully. Pages: {len(doc)}")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text("text")
            full_text += page_text + "\n"
            logger.debug(f"  Page {page_num + 1}: Extracted {len(page_text)} characters")

            # Get embedded links
            links = page.get_links()
            if links:
                logger.debug(f"  Page {page_num + 1}: Found {len(links)} embedded links")
            
            for link in links:
                if link.get('kind') == fitz.LINK_URI:
                    url = link.get('uri', '').strip()
                    if url:
                        rect = link['from']
                        clickable_text = page.get_textbox(rect).strip()
                        link_details.append({
                            'text': clickable_text if clickable_text else url,
                            'url': url
                        })
                        logger.debug(f"    Embedded link: {clickable_text or url}")

        # Get plain-text links
        plain_text_urls = find_urls_in_text(full_text)
        if plain_text_urls:
            logger.debug(f"Found {len(plain_text_urls)} plain-text URLs")
        
        seen_urls = {ld['url'] for ld in link_details}
        for url in plain_text_urls:
            url = url.strip()
            if url not in seen_urls:
                link_details.append({'text': url, 'url': url})
                logger.debug(f"  Added plain-text URL: {url}")

        doc.close()
        logger.debug(f"PDF closed. Total links found: {len(link_details)}")

        # Embed links into the text for the LLM
        sorted_links = sorted(link_details, key=lambda d: len(d['text'] or ''), reverse=True)
        embedded_text = full_text
        links_embedded = 0

        for ld in sorted_links:
            text = ld.get('text') or ''
            url = ld.get('url') or ''
            if not text or not url or text.strip() == url.strip():
                continue

            idx = embedded_text.find(text)
            if idx != -1:
                after_idx = idx + len(text)
                snippet_after = embedded_text[after_idx:after_idx + len(url) + 5]
                if url not in snippet_after:
                    embedded_text = embedded_text[:after_idx] + f" ({url})" + embedded_text[after_idx:]
                    links_embedded += 1

        logger.info(f"✅ PDF parsing complete. Total text length: {len(embedded_text)} characters, Links embedded: {links_embedded}")
        return embedded_text.strip()

    except Exception as e:
        logger.error(f"❌ Error parsing PDF: {e}", exc_info=True)
        return None

def create_docx(resume_data: Dict[str, Any]) -> io.BytesIO:
    """Generates a .docx file from the structured resume JSON."""
    logger.debug("🔧 Creating DOCX file...")
    doc = Document()

    # === SPACING CONFIGURATION ===
    # To change spacing manually, modify these values:
    # HEADER_TO_CONTENT_SPACING = Pt(4)    # Space between section headers and content (increased from Pt(3))
    # INTER_ITEM_SPACING = Pt(2)        # Space between items within sections (decreased from Pt(3.75))
    # INTER_SECTION_SPACING = Pt(5)       # Space between major sections
    # ==============================

    # Set default paragraph spacing to 0 to prevent unwanted gaps
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)

    # --- MARGINS (0.5 inches all sides) ---
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        # Set page size to A4
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # --- CONTACT INFO ---
    contact = resume_data.get('contactInfo', {})
    logger.debug(f"  Adding contact info: {contact.get('name', 'N/A')}")

    # Name: 18pt, bold, Arial, LEFT-ALIGNED
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(contact.get('name', '').replace('\n', ' ').replace('\r', ' '))
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(18)
    run_name.font.bold = True
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.line_spacing = 1.0
    p_name.paragraph_format.space_before = Pt(0)
    # This Pt(5) is the target for spacing BETWEEN items
    p_name.paragraph_format.space_after = Pt(5)

    # Contact Line & Location: 10pt, Arial, LEFT-ALIGNED
    p_contact = doc.add_paragraph()
    
    # Add Contact Line
    run_contact = p_contact.add_run(contact.get('contactLine', '').replace('\n', ' ').replace('\r', ' '))
    run_contact.font.name = 'Arial'
    run_contact.font.size = Pt(10)
    
    # Add Separator
    run_sep = p_contact.add_run(' | ')
    run_sep.font.name = 'Arial'
    run_sep.font.size = Pt(10)
    
    # Add Location
    location_text = contact.get('location', '').replace('Address: ', '').replace('\n', ' ').replace('\r', ' ')
    run_location = p_contact.add_run(location_text) # Get the location field
    run_location.font.name = 'Arial'
    run_location.font.size = Pt(10)
    
    p_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_contact.paragraph_format.line_spacing = 1.0
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(0)

    # Add 10pt spacing after Contact Info
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(10)

    # --- SECTION HEADER HELPER ---
    def add_section_header(title: str):
        """Add a section header: 12pt, NOT bold, uppercase, Arial, LEFT aligned."""
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = False
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)  # HEADER_TO_CONTENT_SPACING - Space between section headers and content

    def add_body_text(text: str):
        """Add body text: 10pt, Arial, LEFT aligned."""
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    def clean_bullet_text(text: str) -> str:
        """Remove common bullet markers from the start of bullet text."""
        text = text.strip()
        # Remove leading bullet markers
        if text.startswith('- ') or text.startswith('• ') or text.startswith('* '):
            text = text[2:].strip()
        elif text.startswith('-') or text.startswith('•') or text.startswith('*'):
            text = text[1:].strip()
        return text

    def add_bullet(text: str):
        """Add bullet point: 10pt, Arial, LEFT aligned."""
        p = doc.add_paragraph(text, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    # --- PROFESSIONAL SUMMARY ---
    logger.debug("  Adding Professional Summary section")
    add_section_header("Professional Summary")

    if resume_data.get('professionalSummary'):
        add_body_text(resume_data.get('professionalSummary', '').replace('\n', ' ').replace('\r', ' '))
    
    # Add 5pt spacing after section
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(5)  # INTER_SECTION_SPACING - Space between Professional Summary and Experience sections

    # --- EXPERIENCE ---
    logger.debug("  Adding Experience section")
    add_section_header("Experience")
    experience_list = resume_data.get('experience', [])
    if experience_list:
        
        for i, exp in enumerate(experience_list):
            logger.debug(f"    Experience {i+1}: {exp.get('title', 'N/A')}")
            
            # NEW STACKED LAYOUT (Fixes Alignment Error)
            p_title = doc.add_paragraph()
            run_title = p_title.add_run(exp.get('title', '').replace('\n', ' ').replace('\r', ' '))
            run_title.font.name = 'Arial'; run_title.font.size = Pt(10); run_title.font.bold = True
            p_title.paragraph_format.space_after = Pt(0); p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            p_company = doc.add_paragraph()
            run_company = p_company.add_run(exp.get('companyLocation', '').replace('\n', ' ').replace('\r', ' '))
            run_company.font.name = 'Arial'; run_company.font.size = Pt(10)
            p_company.add_run(' | ') # Separator
            run_dates = p_company.add_run(exp.get('dates', '').replace('\n', ' ').replace('\r', ' '))
            run_dates.font.name = 'Arial'; run_dates.font.size = Pt(10); run_dates.font.italic = True
            p_company.paragraph_format.space_after = Pt(0); p_company.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Bullets (max 2)
            bullets = exp.get('bullets', [])[:2]
            for j, bullet_text in enumerate(bullets):
                p_bullet = add_bullet(clean_bullet_text(bullet_text))
                if j == len(bullets) - 1 and i < len(experience_list) - 1:
                    p_bullet.paragraph_format.space_after = Pt(2)  # INTER_ITEM_SPACING - Space between experience entries
        
        # Add 5pt spacing after section
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(5)  # INTER_SECTION_SPACING - Space between Experience and Projects sections

    # --- PROJECTS ---
    logger.debug("  Adding Projects section")
    add_section_header("Projects")
    projects_list = resume_data.get('projects', [])
    if projects_list:
        
        for i, proj in enumerate(projects_list):
            logger.debug(f"    Project {i+1}: {proj.get('titleAndTech', 'N/A')}")
            
            # NEW STACKED LAYOUT (Fixes Alignment Error)
            p_title = doc.add_paragraph()
            
            title_and_tech = proj.get('titleAndTech', '').replace('\n', ' ').replace('\r', ' ')
            if ' | Technologies: ' in title_and_tech:
                title_part, tech_part = title_and_tech.split(' | Technologies: ', 1)
                run_title = p_title.add_run(title_part.strip())
                run_title.font.name = 'Arial'; run_title.font.size = Pt(10); run_title.font.bold = True
                run_tech_sep = p_title.add_run(' | Technologies: ')
                run_tech_sep.font.name = 'Arial'; run_tech_sep.font.size = Pt(10); run_tech_sep.font.bold = False
                run_tech = p_title.add_run(tech_part.strip())
                run_tech.font.name = 'Arial'; run_tech.font.size = Pt(10); run_tech.font.bold = False
            else:
                run_title = p_title.add_run(title_and_tech)
                run_title.font.name = 'Arial'; run_title.font.size = Pt(10); run_title.font.bold = True
            
            p_title.paragraph_format.space_after = Pt(0); p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            p_dates = doc.add_paragraph()
            run_dates = p_dates.add_run(proj.get('dates', '').replace('\n', ' ').replace('\r', ' '))
            run_dates.font.name = 'Arial'; run_dates.font.size = Pt(10); run_dates.font.italic = True
            p_dates.paragraph_format.space_after = Pt(0); p_dates.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Bullets (max 2-3)
            bullets = proj.get('bullets', [])[:3]
            for j, bullet_text in enumerate(bullets):
                p_bullet = add_bullet(clean_bullet_text(bullet_text))
                if j == len(bullets) - 1 and i < len(projects_list) - 1:
                    p_bullet.paragraph_format.space_after = Pt(2)  # INTER_ITEM_SPACING - Space between project entries
        
        # Add 5pt spacing after section
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(5)  # INTER_SECTION_SPACING - Space between Projects and Skills sections

    # --- SKILLS ---
    logger.debug("  Adding Skills section (Visible and Invisible)")
    add_section_header("Skills")

    # 1. Add VISIBLE skills (Tier 1 & 2)
    if resume_data.get('skills_visible'):
        add_body_text(resume_data.get('skills_visible', '').replace('\n', ' ').replace('\r', ' '))

    # 2. Add INVISIBLE skills (The Rest)
    if resume_data.get('skills_invisible'):
        p_invisible = doc.add_paragraph()
        run_invisible = p_invisible.add_run(resume_data.get('skills_invisible', '').replace('\n', ' ').replace('\r', ' '))
        run_invisible.font.name = 'Arial'
        run_invisible.font.size = Pt(1)  # Set to 1pt and hidden
        run_invisible.font.hidden = True  # Set font to "hidden"
        p_invisible.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_invisible.paragraph_format.line_spacing = 1.0
        p_invisible.paragraph_format.space_before = Pt(0)
        p_invisible.paragraph_format.space_after = Pt(0)

    # Add 5pt spacing after section
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(5)  # INTER_SECTION_SPACING - Space between Skills and Education sections

    # --- EDUCATION ---
    logger.debug("  Adding Education section")
    add_section_header("Education")
    education_list = resume_data.get('education', [])
    if education_list:
        
        for i, edu in enumerate(education_list):
            logger.debug(f"    Education {i+1}: {edu.get('degree', 'N/A')}")
            
            # Degree, Institution, and Dates on ONE line
            p_degree = doc.add_paragraph()
            
            # Add Degree (NOT bold to reduce bold overuse)
            run_degree = p_degree.add_run(edu.get('degree', '').replace('\n', ' ').replace('\r', ' '))
            run_degree.font.name = 'Arial'
            run_degree.font.size = Pt(10)
            run_degree.font.bold = False
            
            # Add separator (non-bold)
            run_sep = p_degree.add_run(', ')
            run_sep.font.name = 'Arial'
            run_sep.font.size = Pt(10)
            run_sep.font.bold = False # Ensure separator is not bold
            
            # Add Institution & Dates (non-bold)
            run_inst = p_degree.add_run(edu.get('institutionAndDates', '').replace('\n', ' ').replace('\r', ' '))
            run_inst.font.name = 'Arial'
            run_inst.font.size = Pt(10)
            run_inst.font.bold = False

            p_degree.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_degree.paragraph_format.line_spacing = 1.0
            p_degree.paragraph_format.space_before = Pt(0)
            if i < len(education_list) - 1:
                p_degree.paragraph_format.space_after = Pt(2)  # INTER_ITEM_SPACING - Space between education entries
            else:
                p_degree.paragraph_format.space_after = Pt(0)

    # --- SAVE TO BUFFER ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.debug(f"✅ DOCX file created. Size: {len(buffer.getvalue())} bytes")
    return buffer

def create_pdf(resume_data: Dict[str, Any]) -> io.BytesIO:
    """Generates a .pdf file from the structured resume JSON."""
    logger.debug("🔧 Creating PDF file...")
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=0.5*inch, leftMargin=0.5*inch,
        topMargin=0.5*inch, bottomMargin=0.5*inch
    )

    styles = getSampleStyleSheet()
    
    # Define custom styles for consistent formatting
    # Note: Calibri not available in ReportLab; using Helvetica as fallback
    custom_styles = {
        'NameCenter': ParagraphStyle(
            name='NameCenter',
            fontName='Helvetica-Bold',
            fontSize=18,
            alignment=TA_LEFT,  # Changed from TA_CENTER to TA_LEFT
            spaceAfter=5,  # Changed from 0 to 5 to match DOCX spacing between name and contact
            spaceBefore=0,
            leading=18
        ),
        'ContactCenter': ParagraphStyle(
            name='ContactCenter',
            fontName='Helvetica',
            fontSize=10,
            alignment=TA_LEFT,  # Changed from TA_CENTER to TA_LEFT
            spaceAfter=0,
            spaceBefore=0,
            leading=10
        ),
        'SectionTitle': ParagraphStyle(
            name='SectionTitle',
            fontName='Helvetica',  # Changed from Helvetica-Bold to Helvetica
            fontSize=12,
            spaceAfter=4,
            spaceBefore=0,
            leading=12,
            alignment=TA_LEFT
        ),
        'Body': ParagraphStyle(
            name='Body',
            fontName='Helvetica',
            fontSize=10,
            spaceAfter=0,
            spaceBefore=0,
            leading=10,
            alignment=TA_LEFT
        ),
        'Bold10': ParagraphStyle(
            name='Bold10',
            fontName='Helvetica-Bold',
            fontSize=10,
            spaceAfter=0,
            spaceBefore=0,
            leading=10,
            alignment=TA_LEFT
        ),
        'Italic10': ParagraphStyle(
            name='Italic10',
            fontName='Helvetica-Oblique',
            fontSize=10,
            spaceAfter=0,
            spaceBefore=0,
            leading=10,
            alignment=TA_LEFT
        ),
        'Bullet': ParagraphStyle(
            name='Bullet',
            fontName='Helvetica',
            fontSize=10,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=0,
            spaceBefore=0,
            leading=10,
            alignment=TA_LEFT
        ),
        'Invisible': ParagraphStyle(
            name='Invisible',
            fontName='Helvetica',
            fontSize=0.5,  # Extremely small font
            textColor=Color(0, 0, 0, alpha=0.01),  # 1% opaque black (passes contrast check)
            spaceAfter=0,
            spaceBefore=0,
            leading=0.5,
            alignment=TA_LEFT
        )
    }
    
    # Add custom styles to stylesheet
    for style_name, style_obj in custom_styles.items():
        try:
            styles.add(style_obj)
        except:
            pass  # Style may already exist

    def clean_bullet_text(text: str) -> str:
        """Remove common bullet markers from the start of bullet text."""
        text = text.strip()
        # Remove leading bullet markers
        if text.startswith('- ') or text.startswith('• ') or text.startswith('* '):
            text = text[2:].strip()
        elif text.startswith('-') or text.startswith('•') or text.startswith('*'):
            text = text[1:].strip()
        return text

    flowables = []

    # --- CONTACT INFO ---
    contact = resume_data.get('contactInfo', {})
    logger.debug(f"  Adding contact info: {contact.get('name', 'N/A')}")
    
    # Name: 18pt, bold, CENTERED
    name_clean = contact.get('name', '').replace('\n', ' ').replace('\r', ' ')
    flowables.append(Paragraph(name_clean, custom_styles['NameCenter']))
    
    # Contact Line & Location: 10pt, CENTERED
    contact_line_clean = contact.get('contactLine', '').replace('\n', ' ').replace('\r', ' ')
    location_text = contact.get('location', '').replace('Address: ', '').replace('\n', ' ').replace('\r', ' ')
    contact_full = f"{contact_line_clean} | {location_text}"
    flowables.append(Paragraph(contact_full, custom_styles['ContactCenter']))
    
    # Add 10pt spacing after contact (matching DOCX)
    flowables.append(Spacer(1, 10))

    # --- PROFESSIONAL SUMMARY ---
    logger.debug("  Adding Professional Summary section")
    flowables.append(Paragraph("PROFESSIONAL SUMMARY", custom_styles['SectionTitle']))
    # Removed extra spacer after header, spaceAfter=4 on header is sufficient

    if resume_data.get('professionalSummary'):
        summary_clean = resume_data.get('professionalSummary', '').replace('\n', ' ').replace('\r', ' ')
        flowables.append(Paragraph(summary_clean, custom_styles['Body']))
    
    # Add 5pt spacing after section (matching DOCX)
    flowables.append(Spacer(1, 5))

    # --- EXPERIENCE ---
    logger.debug("  Adding Experience section")
    flowables.append(Paragraph("EXPERIENCE", custom_styles['SectionTitle']))
    experience_list = resume_data.get('experience', [])
    if experience_list:
        # Removed extra spacer after header
        
        for i, exp in enumerate(experience_list):
            logger.debug(f"    Experience {i+1}: {exp.get('title', 'N/A')}")
            
            # Job Title: bold 10pt on its own line
            title_clean = exp.get('title', '').replace('\n', ' ').replace('\r', ' ')
            flowables.append(Paragraph(title_clean, custom_styles['Bold10']))
            
            # Company, Location, and Dates on next line with tab
            company_location = exp.get('companyLocation', '').replace('\n', ' ').replace('\r', ' ')
            dates = exp.get('dates', '').replace('\n', ' ').replace('\r', ' ')
            company_dates_line = f"{company_location}\t<i>{dates}</i>"
            flowables.append(Paragraph(company_dates_line, custom_styles['Body']))
            
            # Bullets (max 2)
            for bullet in exp.get('bullets', [])[:2]:
                flowables.append(Paragraph(f"• {clean_bullet_text(bullet)}", custom_styles['Bullet']))
            
            # Add 2pt spacing between entries (matching DOCX)
            if i < len(experience_list) - 1:
                flowables.append(Spacer(1, 2))
        
        # Add 5pt spacing after section (matching DOCX)
        flowables.append(Spacer(1, 5))
    
    # --- PROJECTS ---
    logger.debug("  Adding Projects section")
    flowables.append(Paragraph("PROJECTS", custom_styles['SectionTitle']))
    projects_list = resume_data.get('projects', [])
    if projects_list:
        # Removed extra spacer after header
        
        for i, proj in enumerate(projects_list):
            logger.debug(f"    Project {i+1}: {proj.get('titleAndTech', 'N/A')}")
            
            # Project Name | Technologies AND Dates on one line
            title_and_tech = proj.get('titleAndTech', '').replace('\n', ' ').replace('\r', ' ')
            dates_clean = proj.get('dates', '').replace('\n', ' ').replace('\r', ' ')
            if ' | Technologies: ' in title_and_tech:
                title_part, tech_part = title_and_tech.split(' | Technologies: ', 1)
                project_line = f"<b>{title_part.strip()}</b> | Technologies: {tech_part.strip()}\t<i>{dates_clean}</i>"
            else:
                project_line = f"<b>{title_and_tech}</b>\t<i>{dates_clean}</i>"
            flowables.append(Paragraph(project_line, custom_styles['Body']))
            
            # Bullets (max 2-3)
            for bullet in proj.get('bullets', [])[:3]:
                flowables.append(Paragraph(f"• {clean_bullet_text(bullet)}", custom_styles['Bullet']))
            
            # Add 2pt spacing between entries (matching DOCX)
            if i < len(projects_list) - 1:
                flowables.append(Spacer(1, 2))
        
        # Add 5pt spacing after section (matching DOCX)
        flowables.append(Spacer(1, 5))

    # --- SKILLS ---
    logger.debug("  Adding Skills section (Visible and Invisible)")
    flowables.append(Paragraph("SKILLS", custom_styles['SectionTitle']))

    # 1. Add VISIBLE skills (Tier 1 & 2)
    skills_visible_text = resume_data.get('skills_visible', '').replace('\n', ' ').replace('\r', ' ')
    if skills_visible_text:
        flowables.append(Paragraph(skills_visible_text, custom_styles['Body']))

    # 2. Add INVISIBLE skills (The Rest)
    skills_invisible_text = resume_data.get('skills_invisible', '').replace('\n', ' ').replace('\r', ' ')
    if skills_invisible_text:
        flowables.append(Paragraph(skills_invisible_text, custom_styles['Invisible']))

    # Add 5pt spacing after section (matching DOCX)
    flowables.append(Spacer(1, 5))

    # --- EDUCATION ---
    logger.debug("  Adding Education section")
    flowables.append(Paragraph("EDUCATION", custom_styles['SectionTitle']))
    education_list = resume_data.get('education', [])
    if education_list:
        # Removed extra spacer after header
        
        for i, edu in enumerate(education_list):
            logger.debug(f"    Education {i+1}: {edu.get('degree', 'N/A')}")
            
            # Degree and Institution on one line, Dates tabbed (matching Projects format)
            degree_clean = edu.get('degree', '').replace('\n', ' ').replace('\r', ' ')
            inst_dates_clean = edu.get('institutionAndDates', '').replace('\n', ' ').replace('\r', ' ')
            dates_clean = edu.get('dates', '').replace('\n', ' ').replace('\r', ' ')
            gpa_clean = edu.get('gpa', '').replace('\n', ' ').replace('\r', ' ')
            
            inst_clean = inst_dates_clean.split(',')[0].strip() if ',' in inst_dates_clean else inst_dates_clean
            if gpa_clean:
                education_line = f"{degree_clean}, {inst_clean} | GPA: {gpa_clean}\t<i>{dates_clean}</i>"
            else:
                education_line = f"{degree_clean}, {inst_clean}\t<i>{dates_clean}</i>"
            flowables.append(Paragraph(education_line, custom_styles['Body']))
            
            # Bullets (max 2-3)
            for bullet in edu.get('bullets', [])[:3]:
                flowables.append(Paragraph(f"• {clean_bullet_text(bullet)}", custom_styles['Bullet']))
            
            # Add 2pt spacing between entries (matching DOCX)
            if i < len(education_list) - 1:
                flowables.append(Spacer(1, 2))

    # Build PDF
    doc.build(flowables)
    buffer.seek(0)
    logger.debug(f"✅ PDF file created. Size: {len(buffer.getvalue())} bytes")
    return buffer


# ============================================================================
# PART 6: GEMINI API CALL HELPER (Async)
# ============================================================================

async def call_gemini_api(payload: Dict[str, Any], model: Optional[str] = None) -> Dict[str, Any]:
    """
    Calls the Google Gemini API with exponential backoff.

    If `model` is provided, the request will be sent to that Gemini model.
    Otherwise the configured `GEMINI_API_URL` is used.
    """
    target_model = model if model else None
    if target_model:
        logger.info(f"🔗 Calling Gemini API (model={target_model}) with Google Search Grounding...")
    else:
        logger.info(f"🔗 Calling Gemini API (default) with Google Search Grounding... URL={GEMINI_API_URL}")
    max_retries = 5
    base_delay = 1  # in seconds

    async with httpx.AsyncClient(timeout=120.0) as client:
        for attempt in range(max_retries):
            try:
                if model:
                    full_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={API_KEY}"
                else:
                    full_url = f"{GEMINI_API_URL}?key={API_KEY}"
                logger.debug(f"Attempt {attempt + 1}/{max_retries}: Sending request to Gemini API -> {full_url}")
                
                response = await client.post(
                    full_url,
                    headers={"Content-Type": "application/json"},
                    json=payload
                )

                response.raise_for_status()
                logger.info(f"✅ Gemini API response received successfully (Status: {response.status_code})")
                result = response.json()
                logger.debug(f"Response preview: {str(result)[:500]}...")
                return result

            except httpx.HTTPStatusError as e:
                logger.error(f"❌ HTTP Error: {e.response.status_code}")
                logger.debug(f"Response: {e.response.text}")
                
                if 400 <= e.response.status_code < 500:
                    logger.error(f"Client error (non-retryable): {e.response.status_code}")
                    return {
                        "error": f"Client Error: {e.response.status_code}",
                        "details": e.response.text
                    }

                delay = (base_delay * 2**attempt) + (attempt * 0.1)
                logger.warning(f"⚠️  Server error ({e.response.status_code}). Retrying in {delay:.2f} seconds... (Attempt {attempt + 1}/{max_retries})")
                await asyncio.sleep(delay)

            except httpx.RequestError as e:
                logger.error(f"❌ Request Error: {e}")
                delay = (base_delay * 2**attempt) + (attempt * 0.1)
                logger.warning(f"⚠️  Request error. Retrying in {delay:.2f} seconds... (Attempt {attempt + 1}/{max_retries})")
                await asyncio.sleep(delay)

    logger.error("❌ API call failed after all retry attempts")
    return {"error": "API call failed after several retries."}


# ============================================================================
# PART 7: MAIN FASTAPI ENDPOINT
# ============================================================================

@app.post("/transform-resume/")
async def transform_resume_endpoint(
    resume_file: UploadFile = File(...),
    job_description: str = Form(...),
    target_job_title: str = Form(...),
    time_in_weeks: int = Form(default=2),
    ai_multiplier: int = Form(default=2),
    model: Optional[str] = Form(default=None)
):
    """
    Main endpoint to transform a resume.
    Receives PDF and form data, returns a ZIP file with DOCX, PDF, and condensed JSON response.
    """
    request_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:19]
    logger.info("=" * 80)
    logger.info(f"🚀 NEW RESUME TRANSFORMATION REQUEST (ID: {request_id})")
    logger.info(f"   File: {resume_file.filename}")
    logger.info(f"   Job Description length: {len(job_description)} characters")
    logger.info(f"   Target Job Title: {target_job_title}")
    logger.info(f"   Time in weeks: {time_in_weeks}, AI Multiplier: {ai_multiplier}")
    logger.info("=" * 80)

    # --- 1. Save and Parse PDF ---
    temp_pdf_path = f"temp_{resume_file.filename}"
    try:
        logger.info(f"📥 Saving uploaded file: {temp_pdf_path}")
        # Save uploaded file temporarily
        async with aiofiles.open(temp_pdf_path, 'wb') as out_file:
            content = await resume_file.read()
            await out_file.write(content)
        logger.info(f"✅ File saved. Size: {len(content)} bytes")

        # Parse the saved PDF
        logger.info("📄 Parsing resume PDF...")
        original_resume_text = parse_resume_pdf(temp_pdf_path)
        if not original_resume_text:
            logger.error("❌ Failed to parse PDF - no text extracted")
            raise HTTPException(status_code=400, detail="Could not parse the provided PDF.")
        logger.info(f"✅ PDF parsed successfully. Extracted text: {len(original_resume_text)} characters")

    except HTTPException as e:
        logger.error(f"❌ HTTP Exception during file processing: {e.detail}")
        raise
    except Exception as e:
        logger.error(f"❌ Unexpected error during file processing: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    finally:
        # Clean up the temp file
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)
            logger.debug(f"Cleaned up temporary file: {temp_pdf_path}")

    # --- 2. CALL 1: Analyze Job Description ---
    logger.info("🔨 [CALL 1/5] Constructing JD Analysis prompt...")
    analysis_prompt = JD_ANALYSIS_PROMPT_TEMPLATE.replace(
        "[Insert Job Description Here]", job_description
    )
    logger.debug(f"✅ JD Analysis prompt constructed. Length: {len(analysis_prompt)} characters")

    logger.info("🤖 [CALL 1/5] Calling Gemini API for JD Analysis (JSON Mode)...")
    gemini_payload_call1 = {
        "contents": [
            {"parts": [{"text": analysis_prompt}]}
        ],
        "generationConfig": {"responseMimeType": "application/json"}
    }
    logger.debug(f"Payload size: {len(json.dumps(gemini_payload_call1))} bytes")

    api_response_analysis = await call_gemini_api(gemini_payload_call1, model=model)

    if "error" in api_response_analysis:
        logger.error(f"❌ [CALL 1/5] Gemini API error: {api_response_analysis['error']}")
        raise HTTPException(status_code=500, detail=f"Gemini API Error (JD Analysis): {api_response_analysis['error']}")

    # --- Parse CALL 1 Response ---
    logger.info("📊 [CALL 1/5] Parsing JD Analysis response...")
    try:
        llm_output_text_call1 = api_response_analysis.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
        logger.debug(f"Raw JD Analysis output length: {len(llm_output_text_call1)} characters")

        # Clean up potential markdown code block fences
        if llm_output_text_call1.startswith("```json"):
            logger.debug("Removing markdown JSON fence (start)")
            llm_output_text_call1 = llm_output_text_call1[7:]
        if llm_output_text_call1.endswith("```"):
            logger.debug("Removing markdown JSON fence (end)")
            llm_output_text_call1 = llm_output_text_call1[:-3]

        # Extract JSON from response
        json_match = re.search(r'\{.*\}', llm_output_text_call1, re.DOTALL)
        if json_match:
            llm_output_text_call1 = json_match.group(0)
            logger.debug("Extracted JSON from mixed response")

        logger.debug("Parsing JSON response...")
        jd_analysis_report = json.loads(llm_output_text_call1)
        
        logger.info(f"✅ [CALL 1/5] JD Analysis report generated successfully")
        logger.debug(f"  - Company: {jd_analysis_report.get('company_name', 'N/A')}")
        logger.debug(f"  - Hard Skills: {len(jd_analysis_report.get('keywords', {}).get('hard_skills', []))}")
        logger.debug(f"  - Soft Skills: {len(jd_analysis_report.get('keywords', {}).get('soft_skills', []))}")
        logger.debug(f"  - Tools: {len(jd_analysis_report.get('keywords', {}).get('tools', []))}")
        logger.debug(f"  - Domain Phrases: {len(jd_analysis_report.get('keywords', {}).get('domain_phrases', []))}")

    except (json.JSONDecodeError, IndexError, KeyError, ValueError) as e:
        logger.error(f"❌ [CALL 1/5] Error parsing JD Analysis response: {e}", exc_info=True)
        logger.error(f"Raw JD Analysis Output (first 1000 chars): {llm_output_text_call1[:1000]}")
        raise HTTPException(status_code=500, detail="Failed to parse JD Analysis from the AI. The response may be invalid.")

    # --- 2B. CALL 1.5: Verify Keyword Extraction ---
    logger.info("🔨 [CALL 1.5/5] Constructing Keyword Verification prompt...")

    # Get the list of keywords from the *first* call
    jd_keywords_call1 = jd_analysis_report.get('keywords', {})
    jd_all_skills_list_call1 = (
        jd_keywords_call1.get('hard_skills', []) + jd_keywords_call1.get('tools', []) +
        jd_keywords_call1.get('domain_phrases', []) + jd_keywords_call1.get('soft_skills', [])
    )
    # Get a clean, unique list for the prompt
    jd_all_skills_set_call1 = sorted(list(set(s for s in jd_all_skills_list_call1 if s)))

    verification_prompt = KEYWORD_VERIFICATION_PROMPT_TEMPLATE.replace(
        "[Insert Job Description Here]", job_description
    ).replace(
        "[Insert Keyword List JSON Here]", json.dumps(jd_all_skills_set_call1)
    )
    logger.debug(f"✅ Keyword Verification prompt constructed. Length: {len(verification_prompt)} characters")

    logger.info("🤖 [CALL 1.5/5] Calling Gemini API for Keyword Verification (JSON Mode)...")
    gemini_payload_call1_5 = {
        "contents": [
            {"parts": [{"text": verification_prompt}]}
        ],
        "generationConfig": {"responseMimeType": "application/json"}
    }

    api_response_verification = await call_gemini_api(gemini_payload_call1_5, model=model)

    missing_keywords = []
    if "error" in api_response_verification:
        logger.warning(f"⚠️  [CALL 1.5/5] Gemini API error during verification: {api_response_verification['error']}. Proceeding with original list.")
    else:
        try:
            llm_output_text_call1_5 = api_response_verification.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
            # Clean up potential markdown
            if llm_output_text_call1_5.startswith("```json"):
                llm_output_text_call1_5 = llm_output_text_call1_5[7:]
            if llm_output_text_call1_5.endswith("```"):
                llm_output_text_call1_5 = llm_output_text_call1_5[:-3]
            
            json_match = re.search(r'\{.*\}', llm_output_text_call1_5, re.DOTALL)
            if json_match:
                llm_output_text_call1_5 = json_match.group(0)
                
            verification_data = json.loads(llm_output_text_call1_5)
            missing_keywords = verification_data.get("missing_keywords", [])
            
            if missing_keywords:
                logger.info(f"✅ [CALL 1.5/5] Verification complete. Found {len(missing_keywords)} additional keywords.")
                logger.debug(f"  - Missing: {missing_keywords}")
            else:
                logger.info("✅ [CALL 1.5/5] Verification complete. No missing keywords found.")
                
        except (json.JSONDecodeError, IndexError, KeyError, ValueError) as e:
            logger.warning(f"⚠️  [CALL 1.5/5] Error parsing Verification response: {e}", exc_info=True)
            logger.warning(f"Raw Verification Output: {llm_output_text_call1_5[:500]}")

    # --- Create the new "Master Keyword List" ---
    # This combines the original list + the newly found missing keywords
    master_keyword_list = sorted(list(set(jd_all_skills_set_call1 + missing_keywords)))

    logger.info(f"✨ Master Keyword List created. Total unique keywords: {len(master_keyword_list)}")

    # --- 3. CALL 2: Transform Resume (Text Mode - Main Draft Generator) ---
    logger.info("🔨 [CALL 2/3] Constructing Transform LLM prompt...")
    transform_prompt = RESUME_TRANSFORM_PROMPT_TEMPLATE.replace(
        "[Paste your resume here]", original_resume_text
    ).replace(
        "[Paste job description here]", job_description
    ).replace(
        "[Optional: TIME_IN_WEEKS e.g., 2]", str(time_in_weeks)
    ).replace(
        "[Optional: AI_MULTIPLIER e.g., 3]", str(ai_multiplier)
    )
    logger.debug(f"✅ Transform prompt constructed. Length: {len(transform_prompt)} characters")

    logger.info("🤖 [CALL 2/3] Calling Gemini API for Resume Transformation (Text Mode)...")
    gemini_payload_call2 = {
        "contents": [
            {"parts": [{"text": transform_prompt}]}
        ]
    }
    logger.debug(f"Payload size: {len(json.dumps(gemini_payload_call2))} bytes")

    api_response_transform = await call_gemini_api(gemini_payload_call2, model=model)

    if "error" in api_response_transform:
        logger.error(f"❌ [CALL 2/3] Gemini API error: {api_response_transform['error']}")
        raise HTTPException(status_code=500, detail=f"Gemini API Error (Transform): {api_response_transform['error']}")

    # --- Parse CALL 2 Response ---
    logger.info("📊 [CALL 2/3] Parsing Transform LLM response...")
    try:
        llm_output_text_call2 = api_response_transform.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
        logger.debug(f"Raw Transform output length: {len(llm_output_text_call2)} characters")

        # Clean up potential markdown code block fences
        if llm_output_text_call2.startswith("```json"):
            logger.debug("Removing markdown JSON fence (start)")
            llm_output_text_call2 = llm_output_text_call2[7:]
        if llm_output_text_call2.endswith("```"):
            logger.debug("Removing markdown JSON fence (end)")
            llm_output_text_call2 = llm_output_text_call2[:-3]

        # Extract JSON from response
        json_match = re.search(r'\{.*\}', llm_output_text_call2, re.DOTALL)
        if json_match:
            llm_output_text_call2 = json_match.group(0)
            logger.debug("Extracted JSON from mixed response")

        # Parse the inner JSON
        logger.debug("Parsing JSON response...")
        llm_data_draft = json.loads(llm_output_text_call2)
        logger.debug(f"✅ JSON parsed successfully")

        resume_object_draft = llm_data_draft.get("transformedResume", {}).get("resumeObject", {})
        if not resume_object_draft:
            logger.error("❌ [CALL 2/3] LLM response missing resumeObject")
            raise ValueError("LLM response did not contain a valid resumeObject.")
        
        logger.info(f"✅ [CALL 2/3] Resume object extracted successfully (DRAFT)")
        logger.debug(f"  - Contact: {resume_object_draft.get('contactInfo', {}).get('name', 'N/A')}")
        logger.debug(f"  - Experiences: {len(resume_object_draft.get('experience', []))}")
        logger.debug(f"  - Projects: {len(resume_object_draft.get('projects', []))}")
        logger.debug(f"  - Education: {len(resume_object_draft.get('education', []))}")

    except (json.JSONDecodeError, IndexError, KeyError, ValueError) as e:
        logger.error(f"❌ [CALL 2/3] Error parsing Transform response: {e}", exc_info=True)
        logger.error(f"Raw Transform Output (first 1000 chars): {llm_output_text_call2[:1000]}")
        raise HTTPException(status_code=500, detail="Failed to parse the transformed resume from the AI. The response may be invalid.")

    # --- PRE-PROCESSING FOR CALL 3 ---
    logger.info("🔧 Pre-processing inputs for Call 3 (Final QA)...")
    
    # 1. Get ALL JD keywords as a single string
    jd_keywords = jd_analysis_report.get('keywords', {})
    jd_all_skills_list = (
        jd_keywords.get('hard_skills', []) + jd_keywords.get('tools', []) +
        jd_keywords.get('domain_phrases', []) + jd_keywords.get('soft_skills', [])
    )
    # Ensure uniqueness and preserve rough order
    seen = set()
    jd_all_skills_list_unique = [x for x in jd_all_skills_list if not (x in seen or seen.add(x))]
    jd_all_skills_str = ", ".join(jd_all_skills_list_unique)
    
    # 2. Get the DRAFT resume's CURRENT skills string
    resume_skills_str = resume_object_draft.get('skills', '') # Use the draft from Call 2
    
    logger.debug(f"  - JD All Keywords (string for Call 3): {jd_all_skills_str}")
    logger.debug(f"  - Resume Current Skills (string for Call 3): {resume_skills_str}")

    # --- 4. CALL 3: Final Brute-Force QA (JSON Mode - Keyword Injector + QA) ---
    logger.info("� [CALL 3/3] Constructing Final Brute-Force QA prompt...")
    
    company_name = jd_analysis_report.get('company_name') or 'Target Company'
    
    qa_prompt = FINAL_BRUTEFORCE_QA_PROMPT.replace(
        "[Insert jd_all_skills_str Here]", jd_all_skills_str
    ).replace(
        "[Insert resume_skills_str Here]", resume_skills_str
    ).replace(
        "[Insert Target Job Title String Here]", target_job_title
    ).replace(
        "[Insert Company Name String Here]", company_name
    ).replace(
        "[Insert resumeObject JSON string here]", json.dumps(resume_object_draft)
    )
    logger.debug(f"✅ Final Brute-Force QA prompt constructed. Length: {len(qa_prompt)} characters")
    logger.debug(f"  - Target Job Title: {target_job_title}")
    logger.debug(f"  - Company Name: {company_name}")

    logger.info("🤖 [CALL 3/3] Calling Gemini API for Final Brute-Force QA (JSON Mode)...")
    gemini_payload_call3 = {
        "contents": [
            {"parts": [{"text": qa_prompt}]}
        ],
        "generationConfig": {"responseMimeType": "application/json"}
    }
    logger.debug(f"Payload size: {len(json.dumps(gemini_payload_call3))} bytes")

    api_response_qa = await call_gemini_api(gemini_payload_call3, model=model)

    if "error" in api_response_qa:
        logger.warning(f"⚠️  [CALL 3/3] Gemini API error, using draft version: {api_response_qa['error']}")
        final_resume_object = resume_object_draft
    else:
        # --- Parse CALL 3 Response (JSON Mode: expecting complete resumeObject) ---
        logger.info("📊 [CALL 3/3] Parsing Final Brute-Force QA response...")
        try:
            llm_output_text_call3 = api_response_qa.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
            logger.debug(f"Raw Final QA output length: {len(llm_output_text_call3)} characters")

            # Clean up markdown if present
            llm_output_text_call3 = llm_output_text_call3.strip()
            if llm_output_text_call3.startswith("```json"):
                llm_output_text_call3 = llm_output_text_call3[7:]
            if llm_output_text_call3.startswith("```"):
                llm_output_text_call3 = llm_output_text_call3[3:]
            if llm_output_text_call3.endswith("```"):
                llm_output_text_call3 = llm_output_text_call3[:-3]
            llm_output_text_call3 = llm_output_text_call3.strip()

            # Extract JSON object
            json_match = re.search(r'\{.*\}', llm_output_text_call3, re.DOTALL)
            if json_match:
                llm_output_text_call3 = json_match.group(0)

            logger.debug("Parsing complete resume JSON...")
            final_resume_object = json.loads(llm_output_text_call3)
            
            logger.info(f"✅ [CALL 3/3] Final Brute-Force QA completed successfully")
            logger.debug(f"  - Professional Summary length: {len(final_resume_object.get('professionalSummary', ''))}")
            logger.debug(f"  - Skills length: {len(final_resume_object.get('skills', ''))}")

        except (json.JSONDecodeError, IndexError, KeyError, ValueError) as e:
            logger.warning(f"⚠️  [CALL 3/3] Error parsing Final QA response: {e}", exc_info=True)
            logger.warning(f"Raw Final QA Output (first 500 chars): {llm_output_text_call3[:500]}")
            logger.warning("⚠️  [CALL 3/3] Falling back to draft resume.")
            final_resume_object = resume_object_draft

    # --- 5. Python-Only Cleanup (100% Reliable) ---
    logger.info("🔧 Applying final Python-only cleanup...")

    # Fix 2: Enforce HTTPS on URLs (This logic stays the same)
    if 'contactInfo' in final_resume_object and 'contactLine' in final_resume_object['contactInfo']:
        final_resume_object['contactInfo']['contactLine'] = enforce_absolute_urls(
            final_resume_object['contactInfo']['contactLine']
        )
        logger.info("✅ Absolute HTTPS URLs enforced.")

    # Update the llm_data with final resume object
    final_llm_data = llm_data_draft.copy()
    final_llm_data["transformedResume"]["resumeObject"] = final_resume_object

    # --- 6. Python-Only Brute-Force Keyword Append ---
    logger.info("🔧 [Step 6/6] Applying Python-only brute-force keyword append...")
    try:
        # 1. Get ALL required JD skills
        # Use the "master_keyword_list" we created after the verification call
        logger.info("  Using master keyword list from verification step...")
        jd_skills_set = {skill.lower().strip() for skill in master_keyword_list if skill}
        
        # 2. Get ALL current resume skills
        current_skills_str = final_resume_object.get('skills', '')
        current_skills_list = [skill.strip() for skill in current_skills_str.split(',')]
        # Create a lowercase, stripped set for comparison
        resume_skills_set = {skill.lower().strip() for skill in current_skills_list if skill}

        # 3. Find the missing skills
        missing_skills = jd_skills_set - resume_skills_set
        
        if missing_skills:
            logger.info(f"  Found {len(missing_skills)} missing skills. Appending to skills list...")
            logger.debug(f"  Missing: {missing_skills}")
            
            # 4. Append them
            # We need the *original case* from the master_keyword_list, not the lowercase version
            original_case_missing = [
                skill for skill in set(master_keyword_list) 
                if skill and skill.lower().strip() in missing_skills
            ]
            
            # Combine original list + new list
            final_skills_list = current_skills_list + original_case_missing
            # Ensure no empty strings and reasonable uniqueness
            final_skills_list_cleaned = sorted(list(set(s for s in final_skills_list if s)))
            
            final_resume_object['skills'] = ", ".join(final_skills_list_cleaned)
            logger.info(f"✅ Skills section updated. New length: {len(final_resume_object['skills'])} chars")
        else:
            logger.info("✅ No missing skills found. Skills section is 100% compliant.")

    except Exception as e:
        logger.warning(f"⚠️  Error during Python keyword append: {e}", exc_info=True)
        # Continue with the resume object as-is

    # --- 7. Partition Skills (NEW STEP) ---
    logger.info("📊 Partitioning skills into visible and invisible sections...")
    try:
        # Get the final skills string
        current_skills_str = final_resume_object.get('skills', '')
        
        if current_skills_str:
            # Parse all current skills
            all_skills_list = [skill.strip() for skill in current_skills_str.split(',') if skill.strip()]
            
            # Get Tier 1 (hard_skills) and Tier 2 (soft_skills) from jd_analysis_report
            jd_keywords = jd_analysis_report.get('keywords', {})
            hard_skills = [skill.strip() for skill in jd_keywords.get('hard_skills', []) if skill.strip()]
            soft_skills = [skill.strip() for skill in jd_keywords.get('soft_skills', []) if skill.strip()]
            
            # Create sets for case-insensitive matching
            visible_set = {skill.lower() for skill in hard_skills + soft_skills}
            
            # Partition skills
            visible_skills = []
            invisible_skills = []
            
            for skill in all_skills_list:
                if skill.lower() in visible_set:
                    visible_skills.append(skill)
                else:
                    invisible_skills.append(skill)
            
            # Update final_resume_object with new keys
            final_resume_object['skills_visible'] = ", ".join(visible_skills) if visible_skills else ""
            final_resume_object['skills_invisible'] = ", ".join(invisible_skills) if invisible_skills else ""
            
            logger.info(f"✅ Skills partitioned: {len(visible_skills)} visible, {len(invisible_skills)} invisible")
            logger.debug(f"  Visible (Tier 1 & 2): {final_resume_object['skills_visible']}")
            logger.debug(f"  Invisible (The Rest): {final_resume_object['skills_invisible']}")
        else:
            # No skills found, set empty values
            final_resume_object['skills_visible'] = ""
            final_resume_object['skills_invisible'] = ""
            logger.warning("⚠️  No skills found to partition")
        
        # Clean up the old key
        if 'skills' in final_resume_object:
            del final_resume_object['skills']
            logger.debug("  Deleted old 'skills' key")
            
    except Exception as e:
        logger.warning(f"⚠️  Error during skills partitioning: {e}", exc_info=True)
        # Fallback: keep all skills as visible if partitioning fails
        if 'skills' in final_resume_object:
            final_resume_object['skills_visible'] = final_resume_object['skills']
            final_resume_object['skills_invisible'] = ""
            del final_resume_object['skills']

    # --- 8. Generate Output Files ---
    logger.info("📝 Generating output documents...")
    try:
        logger.debug("Generating DOCX...")
        docx_buffer = create_docx(final_resume_object)
        
        logger.debug("Generating PDF...")
        pdf_buffer = create_pdf(final_resume_object)
        logger.info(f"✅ Documents generated successfully")
    except Exception as e:
        logger.error(f"❌ Error generating documents: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate output documents: {str(e)}")

    # --- 6. Return DOCX File Directly ---
    base_filename = final_llm_data.get("transformedResume", {}).get("suggestedFilename", "resume").replace(".pdf", "")
    logger.debug(f"Suggested filename: {base_filename}")
    
    logger.info(f"✅ RESUME TRANSFORMATION COMPLETE (ID: {request_id})")
    logger.info(f"   3-Call Pipeline: JD Analysis → Transform → Keyword Injection → Final QA")
    logger.info(f"   Returning DOCX: {base_filename}.docx")
    logger.info("=" * 80)
    
    return Response(
        content=docx_buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={base_filename}.docx"
        }
    )


@app.get("/health")
async def health_check():
    """Health check endpoint."""
    logger.debug("Health check requested")
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}


# ============================================================================
# PART 8: UVICORN RUNNER
# ============================================================================

if __name__ == "__main__":
    logger.info("🚀 Starting Uvicorn server...")
    logger.info(f"📍 API will be available at: http://0.0.0.0:8000")
    logger.info(f"📚 API documentation at: http://0.0.0.0:8000/docs")
    
    # Suppress uvicorn access logs
    logging.getLogger("uvicorn.access").setLevel(logging.WARNING)
    
    uvicorn.run(app, host="0.0.0.0", port=8000, access_log=False)
